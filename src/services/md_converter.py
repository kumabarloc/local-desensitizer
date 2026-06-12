"""Markdown 中间格式转换器

提供 docx ↔ Markdown 的双向转换，作为脱敏流程的中间格式。
- docx → MD：委托给 DocxHandler.read()（使用 python-docx 按 body 顺序遍历，
            不走 mammoth——mammoth 转 MD 会丢失表格结构）
- MD → docx：自实现解析器，支持表格/标题/列表/代码块
"""
from pathlib import Path


def docx_to_markdown(path: Path) -> tuple[str, list[str]]:
    """docx → Markdown

    委托给 DocxHandler.read()，保留表格/标题/列表结构。
    Returns:
        (markdown_text, warnings)  warnings 预留，当前为空
    """
    from src.services.document_handler import DocxHandler
    handler = DocxHandler()
    md_text = handler.read(path)
    return md_text, []


def markdown_to_docx(md_text: str, path: Path) -> None:
    """Markdown → docx

    支持以下 MD 元素：
    - 标题：# / ## / ### / ####
    - 列表：- / * / 1.
    - 表格：| col1 | col2 |\n| --- | --- |\n| a | b |
    - 代码块：```
    - 普通段落
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    # 全文样式微调（保持中文可读）
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    lines = md_text.split('\n')
    i = 0
    in_code_block = False
    code_buffer: list[str] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 代码块（围栏）: ```
        if stripped.startswith('```'):
            if in_code_block:
                # 结束
                _write_code_block(doc, code_buffer)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # 表格：以 | 开头且包含 |
        if _is_table_row(stripped):
            block = []
            while i < len(lines) and _is_table_row(lines[i].strip()):
                block.append(lines[i])
                i += 1
            _write_table(doc, block)
            continue

        # 标题
        if _is_heading(stripped):
            level = _heading_level(stripped)
            text = stripped[level:].strip()
            doc.add_heading(text, level=min(level, 6))
            i += 1
            continue

        # 无序列表
        if _is_unordered_list_item(stripped):
            text = _strip_list_marker(stripped, ordered=False)
            doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue

        # 有序列表
        if _is_ordered_list_item(stripped):
            text = _strip_list_marker(stripped, ordered=True)
            doc.add_paragraph(text, style='List Number')
            i += 1
            continue

        # 空行：跳过
        if not stripped:
            i += 1
            continue

        # 普通段落
        doc.add_paragraph(stripped)
        i += 1

    doc.save(path)


# ============ 内部辅助函数 ============

def _is_table_row(s: str) -> bool:
    """判断是否是表格行：以 | 开头、以 | 结尾（或仅含 | 分隔）"""
    return s.startswith('|') and '|' in s[1:]


def _is_heading(s: str) -> bool:
    """判断是否是 MD 标题。允许 1-6 个 # 开头，后跟空格或字符串结束。"""
    if not s.startswith('#'):
        return False
    i = 0
    while i < len(s) and s[i] == '#':
        i += 1
    if i == 0:
        return False
    if i > 6:
        return False
    if i >= len(s):
        return True
    return s[i] in ' \t'


def _heading_level(s: str) -> int:
    """返回 # 的数量（1-6），超过 6 当作 6"""
    level = 0
    for ch in s:
        if ch == '#':
            level += 1
        else:
            break
    return min(level, 6)


def _is_unordered_list_item(s: str) -> bool:
    return len(s) >= 2 and s[0] in ('-', '*', '+') and s[1] in (' ', '\t')


def _is_ordered_list_item(s: str) -> bool:
    """形如 1. / 23. 这样的有序列表"""
    if len(s) < 3:
        return False
    digits = ''
    for ch in s:
        if ch.isdigit():
            digits += ch
        else:
            break
    return bool(digits) and len(s) > len(digits) and s[len(digits)] == '.' and s[len(digits) + 1] in (' ', '\t')


def _strip_list_marker(s: str, ordered: bool) -> str:
    if ordered:
        # 跳过 数字+.+空格
        idx = 0
        while idx < len(s) and s[idx].isdigit():
            idx += 1
        return s[idx + 2:].strip()
    else:
        return s[2:].strip()


def _write_table(doc, block: list[str]) -> None:
    """写 docx 表格。block 是一组 |col|col| 行，第二行通常为 ---|---|--- 分隔。"""
    if not block:
        return
    rows = []
    for line in block:
        # 去掉首尾 | 再按 | 切
        inner = line.strip()
        if inner.startswith('|'):
            inner = inner[1:]
        if inner.endswith('|'):
            inner = inner[:-1]
        cells = [c.strip() for c in inner.split('|')]
        rows.append(cells)

    if not rows:
        return

    # 检测分隔行（---|---|---）
    if len(rows) >= 2 and all(_is_separator_cell(c) for c in rows[1]):
        header = rows[0]
        body = rows[2:]
    else:
        header = rows[0]
        body = rows[1:] if len(rows) > 1 else []

    if not header:
        return
    n_cols = len(header)
    table = doc.add_table(rows=1 + len(body), cols=n_cols)
    try:
        table.style = 'Table Grid'
    except KeyError:
        pass  # 某些环境没有 Table Grid style

    for c, text in enumerate(header):
        table.rows[0].cells[c].text = text
    for r, row in enumerate(body, start=1):
        for c, text in enumerate(row):
            if c < n_cols:
                table.rows[r].cells[c].text = text


def _is_separator_cell(s: str) -> bool:
    """判断是否是表格分隔单元格（--- 或 :---: 等）"""
    s = s.strip()
    if not s:
        return False
    cleaned = s.replace(':', '').replace('-', '').strip()
    return not cleaned  # 只含 - 和 :


def _write_code_block(doc, code_lines: list[str]) -> None:
    """写代码块（用等宽字体段落）"""
    from docx.shared import Pt
    if not code_lines:
        return
    p = doc.add_paragraph()
    run = p.add_run('\n'.join(code_lines))
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
