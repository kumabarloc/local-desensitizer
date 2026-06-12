"""MD 中间格式转换器测试"""
import pytest
from pathlib import Path
import tempfile

from src.services.md_converter import (
    docx_to_markdown,
    markdown_to_docx,
    _is_table_row,
    _is_heading,
    _heading_level,
    _is_unordered_list_item,
    _is_ordered_list_item,
    _is_separator_cell,
)


class TestHelpers:
    """内部辅助函数测试"""

    def test_table_row_detection(self):
        assert _is_table_row("| col1 | col2 |")
        assert _is_table_row("| --- | --- |")
        assert _is_table_row("| a | b | c | d |")
        assert not _is_table_row("普通段落")
        assert not _is_table_row("| 单独一个管道")
        assert not _is_table_row("")

    def test_heading_detection(self):
        assert _is_heading("# 一级标题")
        assert _is_heading("## 二级标题")
        assert _is_heading("###### 六级标题")
        assert _is_heading("#")
        assert not _is_heading("普通段落")
        assert not _is_heading("#hashtag 不是标题")  # 标题后必须有空格

    def test_heading_level(self):
        assert _heading_level("# h1") == 1
        assert _heading_level("## h2") == 2
        assert _heading_level("###### h6") == 6
        assert _heading_level("####### h7") == 6  # 超过 6 截断

    def test_unordered_list(self):
        assert _is_unordered_list_item("- 项目")
        assert _is_unordered_list_item("* 项目")
        assert _is_unordered_list_item("+ 项目")
        assert not _is_unordered_list_item("普通段落")
        assert not _is_unordered_list_item("-缺空格")  # 标记后必须有空格

    def test_ordered_list(self):
        assert _is_ordered_list_item("1. 第一")
        assert _is_ordered_list_item("23. 第二十三")
        assert not _is_ordered_list_item("普通段落")
        assert not _is_ordered_list_item("1.缺空格")

    def test_separator_cell(self):
        assert _is_separator_cell("---")
        assert _is_separator_cell(":---:")
        assert _is_separator_cell(":---")
        assert _is_separator_cell("---:")
        assert not _is_separator_cell("数据")
        assert not _is_separator_cell("")


class TestMarkdownToDocx:
    """MD → docx 转换测试"""

    def test_simple_paragraph(self, tmp_path):
        md_text = "这是第一段。\n\n这是第二段。"
        out = tmp_path / "out.docx"
        markdown_to_docx(md_text, out)
        assert out.exists()
        assert out.stat().st_size > 0

    def test_headings(self, tmp_path):
        md_text = "# 一级\n\n## 二级\n\n### 三级"
        out = tmp_path / "out.docx"
        markdown_to_docx(md_text, out)
        assert out.exists()

    def test_lists(self, tmp_path):
        md_text = "- 项目1\n- 项目2\n- 项目3\n\n1. 步骤1\n2. 步骤2"
        out = tmp_path / "out.docx"
        markdown_to_docx(md_text, out)
        assert out.exists()

    def test_simple_table(self, tmp_path):
        """这是核心测试：表格在 docx 里要存在"""
        from docx import Document
        md_text = """# 报告

| 姓名 | 部门 | 金额 |
| --- | --- | --- |
| 张三 | 监测室 | 500万 |
| 李四 | 数据中心 | 1000万 |

## 结论

数据如上。"""
        out = tmp_path / "out.docx"
        markdown_to_docx(md_text, out)
        assert out.exists()

        # 验证：读回来应该有表格
        doc = Document(out)
        assert len(doc.tables) >= 1, "应该有至少 1 个表格"
        table = doc.tables[0]
        assert len(table.rows) == 3, "表格应该有 3 行（1 表头 + 2 数据）"
        assert len(table.columns) == 3, "表格应该有 3 列"
        assert table.rows[0].cells[0].text == "姓名"
        assert table.rows[1].cells[0].text == "张三"
        assert table.rows[1].cells[2].text == "500万"
        assert table.rows[2].cells[1].text == "数据中心"

    def test_table_no_separator(self, tmp_path):
        """没有分隔行的表格也要能处理"""
        from docx import Document
        md_text = """| A | B |
| 1 | 2 |"""
        out = tmp_path / "out.docx"
        markdown_to_docx(md_text, out)
        doc = Document(out)
        assert len(doc.tables) == 1
        assert doc.tables[0].rows[0].cells[0].text == "A"

    def test_code_block(self, tmp_path):
        md_text = "```python\ndef hello():\n    print('hi')\n```"
        out = tmp_path / "out.docx"
        markdown_to_docx(md_text, out)
        assert out.exists()


class TestDocxToMarkdown:
    """docx → MD 转换测试"""

    def test_simple_docx(self, tmp_path):
        from docx import Document
        doc = Document()
        doc.add_paragraph("第一段")
        doc.add_paragraph("第二段")
        src = tmp_path / "src.docx"
        doc.save(src)

        md_text, warnings = docx_to_markdown(src)
        assert "第一段" in md_text
        assert "第二段" in md_text

    def test_docx_with_table(self, tmp_path):
        """核心测试：docx 表格要保留到 MD"""
        from docx import Document
        doc = Document()
        doc.add_paragraph("报告标题")
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        table.rows[0].cells[0].text = "姓名"
        table.rows[0].cells[1].text = "金额"
        table.rows[1].cells[0].text = "李四"
        table.rows[1].cells[1].text = "500万"
        src = tmp_path / "src.docx"
        doc.save(src)

        md_text, warnings = docx_to_markdown(src)
        # mammoth 输出的 MD 表格用 | 语法
        assert "姓名" in md_text
        assert "金额" in md_text
        assert "李四" in md_text
        assert "500万" in md_text
        # 至少有一行表格
        assert "|" in md_text


class TestDocxRoundTrip:
    """docx → MD → docx 完整往返"""

    def test_table_round_trip(self, tmp_path):
        """核心测试：表格能完整往返"""
        from docx import Document
        # 准备源 docx
        doc = Document()
        doc.add_paragraph("报告")
        table = doc.add_table(rows=3, cols=3)
        table.style = 'Table Grid'
        cells_data = [
            ["姓名", "部门", "金额"],
            ["张三", "监测室", "500万"],
            ["李四", "数据中心", "1000万"],
        ]
        for r, row in enumerate(cells_data):
            for c, text in enumerate(row):
                table.rows[r].cells[c].text = text
        src = tmp_path / "src.docx"
        doc.save(src)

        # docx → MD
        md_text, _ = docx_to_markdown(src)

        # MD → docx
        out = tmp_path / "out.docx"
        markdown_to_docx(md_text, out)

        # 验证
        result = Document(out)
        assert len(result.tables) == 1, "表格应该在往返后保留"
        result_table = result.tables[0]
        assert len(result_table.rows) == 3
        assert len(result_table.columns) == 3
        assert result_table.rows[0].cells[0].text == "姓名"
        assert result_table.rows[1].cells[0].text == "张三"
        assert result_table.rows[2].cells[1].text == "数据中心"
        assert result_table.rows[2].cells[2].text == "1000万"
